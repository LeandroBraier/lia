# -*- coding: utf-8 -*-
"""
 * @license
 * SPDX-License-Identifier: Apache-2.0
 """

import os
import re
import json
import shutil
import pandas as pd
from docx import Document
import openpyxl

# Componentes de Microsoft Presidio
try:
    from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
    from presidio_analyzer.nlp_engine import NlpEngineProvider
    from presidio_anonymizer import AnonymizerEngine
    from presidio_anonymizer.entities import OperatorConfig
except ImportError:
    class NlpEngineProvider:
        def __init__(self, **kwargs): pass
        def create_engine(self): return None
    class Pattern:
        def __init__(self, **kwargs): pass
    class PatternRecognizer:
        def __init__(self, **kwargs): pass
    class AnalyzerEngine:
        def __init__(self, **kwargs): pass
        def analyze(self, **kwargs): return []
    class AnonymizerEngine:
        def anonymize(self, **kwargs):
            class Res: text = kwargs.get('text', '')
            return Res()

# Parche SSL para descarga de modelos en macOS
import ssl
try:
    ssl._create_default_https_context = ssl._create_unverified_context
except AttributeError:
    pass

# Soporte PDF e Imágenes (OCR)
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import easyocr
    import cv2
    from PIL import Image
except ImportError:
    easyocr = None
    cv2 = None
    Image = None

# 1. Configuración de rutas corporativas
# Detectar si estamos en un sistema de solo lectura (ej: DMG montado en macOS)
_base_dir = os.path.dirname(os.path.abspath(__file__))
if not os.access(_base_dir, os.W_OK) or _base_dir.startswith("/Volumes/"):
    # Sistema de solo lectura: usar carpeta en el Home del usuario
    _base_dir = os.path.join(os.path.expanduser("~"), "LiaVault")

CARPETA_ENTRADA = os.path.join(_base_dir, "entrada")
CARPETA_SALIDA_DEFECTO = os.path.join(_base_dir, "Archivos seguros de Checkpoint Ley IA")
CARPETA_SALIDA = CARPETA_SALIDA_DEFECTO
CARPETA_PROCESADOS = os.path.join(_base_dir, "procesados")
CARPETA_CONFIG = os.path.join(_base_dir, "config")
RUTA_DICCIONARIO = os.path.join(CARPETA_CONFIG, "diccionario_exclusiones.txt")

try:
    os.makedirs(CARPETA_ENTRADA, exist_ok=True)
    os.makedirs(CARPETA_SALIDA_DEFECTO, exist_ok=True)
    os.makedirs(CARPETA_PROCESADOS, exist_ok=True)
    os.makedirs(CARPETA_CONFIG, exist_ok=True)
except Exception as ex_mk:
    print(f"⚠️ Aviso al preparar directorios: {ex_mk}")

    if not os.path.exists(RUTA_DICCIONARIO):
        with open(RUTA_DICCIONARIO, "w", encoding="utf-8") as f:
            f.write("# Lista de palabras prohibidas de la empresa (un término por línea)\n")
            f.write("# Las coincidencias se reemplazarán por [TERMINO_CONFIDENCIAL]\n")
            f.write("Proyecto Halcon\n")
            f.write("Adquisición Alfa\n")
            f.write("Formula-X7\n")
except OSError:
    pass


MAPA_ETIQUETAS = {
    "PERSON": "[PERSONA]",
    "LOCATION": "[UBICACION]",
    "ORGANIZATION": "[ORGANIZACION]",
    "EMAIL_ADDRESS": "[EMAIL]",
    "PHONE_NUMBER": "[TELEFONO]",
    "CREDIT_CARD": "[TARJETA]"
}

# 2. Inicializar IA local y Lector OCR
print("⏳ Cargando modelos e IA local de Lia Vault...")
try:
    nlp_config = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "es", "model_name": "es_core_news_sm"}]
    }
    provider = NlpEngineProvider(nlp_configuration=nlp_config)
    nlp_engine = provider.create_engine()
    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["es"]) if nlp_engine else AnalyzerEngine(supported_languages=["es"])
    anonymizer = AnonymizerEngine()
    
    patron_tarjeta = Pattern(name="tarjeta_patron", regex=r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b", score=0.85)
    reconocedor_tarjetas = PatternRecognizer(supported_entity="CREDIT_CARD", supported_language="es", patterns=[patron_tarjeta])
    analyzer.registry.add_recognizer(reconocedor_tarjetas)
except Exception as e:
    print(f"⚠️ Alerta: Presidio o spaCy no cargados completamente: {e}")
    analyzer = None
    anonymizer = None

try:
    if easyocr:
        lector_ocr = easyocr.Reader(['es', 'en'], gpu=False)
    else:
        lector_ocr = None
except Exception as e:
    print(f"⚠️ Alerta: EasyOCR no inicializado: {e}")
    lector_ocr = None


def cargar_diccionario_corporativo():
    terminos = []
    if os.path.exists(RUTA_DICCIONARIO):
        with open(RUTA_DICCIONARIO, "r", encoding="utf-8") as f:
            for linea in f:
                linea = linea.strip()
                if linea and not linea.startswith("#"):
                    terminos.append(linea)
    return terminos


SOFTWARE_ALLOWLIST = {
    'chatgpt', 'copilot', 'gemini', 'claude', 'llama', 'gpt', 'gpt-4', 'gpt-3.5',
    'windows', 'office', 'excel', 'word', 'powerpoint', 'google', 'microsoft',
    'python', 'docker', 'react', 'vite', 'node', 'express', 'spacy', 'presidio',
    'easyocr', 'pdf', 'csv', 'txt', 'docx', 'xlsx', 'json', 'html', 'css', 'javascript',
    'typescript', 'flet', 'github', 'gitlab', 'slack', 'teams', 'zoom',
    'meeting', 'inteligencia', 'artificial', 'bueno', 'buenos', 'días', 'hola',
    'gracias', 'perdón', 'claro', 'exacto', 'así', 'igual', 'cualquier', 'esa',
    'uno', 'vale', 'ah', 'mirá', 'mira', 'acá', 'después', 'grabación',
    'aparentemente', 'socios', 'pymes', 'tecnologías', 'verdadero', 'consciente',
    'julio', 'agosto', 'unión europea', 'barcelona', 'view', 'notes', 'cuál', 'quién',
    # Términos genéricos corporativos / estadísticos
    'empresa', 'empresas', 'compañía', 'compania', 'entidad', 'cliente', 'clientes',
    'sociedad', 'director', 'directora', 'gerente', 'jefe', 'jefa', 'empleado', 'empleada',
    'usuario', 'sistema', 'servicio', 'data', 'año', 'unidad', 'número', 'valor', 'concepto',
    'estado', 'tipo', 'hombres', 'mujeres', 'alumnos', 'alumnas', 'matriculados', 'matriculadas',
    'total', 'otros', 'provisional', 'definitivo', 'territorio', 'código',
    'primer', 'segundo', 'curso', 'ciclo', 'formación', 'profesional', 'básica',
    'centros', 'escolares', 'formativos', 'dato', 'nacional',
    'rol', 'ajuste', 'score', 'rating', 'icp', 'lookalikes', 'lista',
    # Términos legales y estructurales de NDAs para evitar falsos positivos
    'apoderado', 'apoderada', 'titular', 'presidente', 'presidenta', 'vicepresidente', 
    'vicepresidenta', 'secretario', 'secretaria', 'vocal', 'administrador', 'administradora', 
    'socio', 'socia', 'contratante', 'proveedora', 'receptora', 'divulgadora', 'parte', 'partes', 
    'objeto', 'definiciones', 'excepciones', 'devolución', 'destrucción', 'incumplimiento', 
    'duración', 'jurisdicción', 'anexo', 'anexos', 'acuerdo', 'contrato', 'convenio', 'reunidos', 
    'exclusiones', 'primera', 'segunda', 'tercera', 'cuarta', 'quinta', 'sexta', 'séptima', 
    'octava', 'novena', 'décima', 'primero', 'segundo', 'tercero', 'cuarto', 'quinto', 'sexto', 
    'séptimo', 'octavo', 'noveno', 'décimo', 'información', 'confidencial', 'confidencialidad'
}

SPANISH_NAMES = [
    "Carlos Mendoza", "Santiago", "Alejandro Gómez", "Juan Pérez", "María Rodríguez",
    "Laura Martínez", "Andrés Felipe", "Sofía Castro", "Santiago Valencia", "Gabriela",
    "Diego Armando", "Carlos", "Juan", "Gómez", "Mendoza", "Alejandro", "Sonia", "Beatriz",
    "Laura", "Noelia", "Aznar", "Pedro", "Ana", "Lucía", "Javier", "Martín", "Elena", "Carmen"
]

def aplicar_diccionario_exclusiones(texto):
    if not isinstance(texto, str):
        return texto
    terminos = cargar_diccionario_corporativo()
    for termino in terminos:
        patron = re.compile(re.escape(termino), re.IGNORECASE)
        texto = patron.sub("[TERMINO_CONFIDENCIAL]", texto)
    return texto

MAPA_ETIQUETAS = {
    "PERSON": "PERSONA",
    "LOCATION": "UBICACION",
    "ORGANIZATION": "ORGANIZACION",
    "EMAIL_ADDRESS": "CORREO",
    "PHONE_NUMBER": "TELEFONO",
    "CREDIT_CARD": "TARJETA_CREDITO",
    "DOMAIN": "DOMINIO",
    "SECRET_KEY": "CLAVE",
    "USERNAME": "USUARIO"
}

def sanitizar_nombre_archivo(nombre_archivo, carpeta_salida=None):
    nombre_base, ext = os.path.splitext(nombre_archivo)
    nombre_base_espacios = nombre_base.replace("_", " ").replace("-", " ")
    limpio, key_map = anonimizar_texto_con_ia_y_key(nombre_base_espacios, carpeta_salida=carpeta_salida)
    sanitizado_base = limpio.replace(" ", "_")
    return f"{sanitizado_base}{ext}", key_map

def all_caps_to_title_case(text):
    exclusions = {'DNI', 'NIE', 'NIF', 'IVA', 'NDA'}
    def replace_word(match):
        word = match.group(0)
        if word in exclusions:
            return word
        return word.title()
    return re.sub(r'\b[A-ZÁÉÍÓÚÑ]{2,}\b', replace_word, text)

def anonimizar_texto_con_ia_y_key(texto, nombre_archivo=None, carpeta_salida=None, contadores_existentes=None, key_map_existente=None):
    if not isinstance(texto, str) or not texto.strip():
        return texto, {}

    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    os.makedirs(target_salida, exist_ok=True)
    texto_filtrado = aplicar_diccionario_exclusiones(texto)

    key_map = key_map_existente if key_map_existente is not None else {}
    contadores = contadores_existentes if contadores_existentes is not None else {}
    texto_anom = texto_filtrado

    valor_a_etiqueta = {val: k for k, val in key_map.items()}

    # 1. Presidio AI Engine (si está disponible)
    if analyzer:
        try:
            # Preprocesar mayúsculas sostenidas a Title Case para asegurar detección por spaCy
            texto_analisis = all_caps_to_title_case(texto_anom)
            resultados = analyzer.analyze(
                text=texto_analisis,
                language="es",
                entities=["PERSON", "LOCATION", "ORGANIZATION", "EMAIL_ADDRESS", "PHONE_NUMBER", "CREDIT_CARD"]
            )
            for pii in sorted(resultados, key=lambda x: x.start, reverse=True):
                # Recuperar el valor original del texto con las mayúsculas originales
                val = texto_anom[pii.start:pii.end].strip()
                if not val or val.lower() in SOFTWARE_ALLOWLIST or len(val) < 2:
                    continue
                tag = MAPA_ETIQUETAS.get(pii.entity_type, "PERSONA")
                if val in valor_a_etiqueta:
                    etiqueta = valor_a_etiqueta[val]
                else:
                    idx = contadores.get(tag, 1)
                    etiqueta = f"[{tag}_{idx}]"
                    contadores[tag] = idx + 1
                    key_map[etiqueta] = val
                    valor_a_etiqueta[val] = etiqueta

                # Reemplazar únicamente en la posición detectada para evitar falsos positivos
                texto_anom = texto_anom[:pii.start] + etiqueta + texto_anom[pii.end:]
        except Exception as ex:
            print(f"Error analizando Presidio: {ex}")

    # Actualizar mapa de reversión después del procesamiento de Presidio
    valor_a_etiqueta = {val: k for k, val in key_map.items()}

    # 2. Expresiones Regulares Específicas
    # Las expresiones de vocativos y tags de interlocutores son sensibles a mayúsculas para evitar clasificar preposiciones/verbos
    PATRONES_REGEX = [
        (re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", re.IGNORECASE), "CORREO"),
        (re.compile(r"\b(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-]+)*\.(?:com|es|org|net|co|io|gov|edu))\b", re.IGNORECASE), "DOMINIO"),
        (re.compile(r"\b(?:sk-[a-zA-Z0-9]{20,}|ghp_[a-zA-Z0-9]{20,}|bearer\s+[a-zA-Z0-9._\-]+|(?:password|clave|contraseña|pwd|api_key|secret|token)\s*[:=]\s*['\"]?([^\s'\"]{4,})['\"]?)\b", re.IGNORECASE), "CLAVE"),
        (re.compile(r"(?:@([a-zA-Z0-9._-]+)|\b(?:user|usuario|username|interlocutor)\s*[:=]\s*['\"]?([a-zA-Z0-9._-]+)['\"]?)", re.IGNORECASE), "USUARIO"),
        (re.compile(r"\b(?:\d[ -]?){13,16}\b"), "TARJETA_CREDITO"),
        (re.compile(r"\+?\b\d{1,4}[-\s]?\d{3,4}[-\s]?\d{3,4}\b"), "TELEFONO"),
        (re.compile(r"\b\d{7,10}[-\s]?[A-Z0-9]?\b"), "ID_OFICIAL"),
        # Nombres en mayúsculas sostenidas (ALL CAPS) y nombres completos tipo Title Case
        (re.compile(r"\b[A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})+\b"), "PERSONA"),
        (re.compile(r"\b[A-ZÁÉÍÓÚÑ][a-záéíóúñ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]{2,})+\b"), "PERSONA"),
        # Fórmulas de introducción legal (ej: representada por, en este acto por) para capturar nombres incluso cortados
        (re.compile(r"\b(?:representado|representada|firmado|firma|por)\s+(?:en\s+este\s+acto\s+)?(?:por\s+)?([A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})*|[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)*)\b"), "PERSONA"),
        (re.compile(r"(?:(?:\d{2}:\d{2}\s+)?([A-ZÁÉÍÓÚÑ][a-záéíóúñ]{2,}(?:\s+[A-Z]\.?[A-Z]\.?)?)\s*:)"), "PERSONA"),
        (re.compile(r"(?:,\s*|\b(?:gracias|mira|mirá|dime|decía|bueno|hola|estimado|estimada|saludos)\s+)([A-ZÁÉÍÓÚÑ][a-záéíóúñ]{2,})"), "PERSONA"),
    ]

    # Matching explícito de nombres propios en español
    for nombre in SPANISH_NAMES:
        if nombre.lower() in SOFTWARE_ALLOWLIST:
            continue
        patron = re.compile(rf"\b{re.escape(nombre)}\b", re.IGNORECASE)
        matches = list(patron.finditer(texto_anom))
        for match in matches:
            val = match.group(0)
            if val.startswith("[") and val.endswith("]"):
                continue
            if val in valor_a_etiqueta:
                etiqueta = valor_a_etiqueta[val]
            else:
                idx = contadores.get("PERSONA", 1)
                etiqueta = f"[PERSONA_{idx}]"
                contadores["PERSONA"] = idx + 1
                key_map[etiqueta] = val
                valor_a_etiqueta[val] = etiqueta
            texto_anom = re.sub(rf"\b{re.escape(val)}\b", etiqueta, texto_anom)

    for regex, tag in PATRONES_REGEX:
        matches = regex.finditer(texto_anom)
        for match in list(matches):
            val = match.group(1) if match.groups() and match.group(1) else match.group(0)
            if not val or (val.startswith("[") and val.endswith("]")):
                continue
            if val.lower().strip() in SOFTWARE_ALLOWLIST:
                continue
            
            if val in valor_a_etiqueta:
                etiqueta = valor_a_etiqueta[val]
            else:
                idx = contadores.get(tag, 1)
                etiqueta = f"[{tag}_{idx}]"
                contadores[tag] = idx + 1
                key_map[etiqueta] = val
                valor_a_etiqueta[val] = etiqueta

            # Reemplazar con límites de palabra para tipos textuales y exacto para secuencias/datos
            if tag in ["PERSONA", "ORGANIZACION", "UBICACION", "ID_OFICIAL"]:
                texto_anom = re.sub(rf"\b{re.escape(val)}\b", etiqueta, texto_anom)
            else:
                texto_anom = texto_anom.replace(val, etiqueta)

    if nombre_archivo and key_map:
        guardar_llave_key(target_salida, nombre_archivo, key_map)

    return texto_anom, key_map


def revertir_anonimizacion(texto_anonimizado, mapa_reversion):
    if not isinstance(texto_anonimizado, str) or not mapa_reversion:
        return texto_anonimizado
    
    texto_restaurado = texto_anonimizado
    for etiqueta, original in mapa_reversion.items():
        if not etiqueta or not isinstance(etiqueta, str):
            continue
        etiqueta_limpia = etiqueta.strip("[]")
        texto_restaurado = re.sub(r"\[+" + re.escape(etiqueta_limpia) + r"\]+", str(original), texto_restaurado)
        texto_restaurado = texto_restaurado.replace(etiqueta, str(original))
    return texto_restaurado


def ejecutar_reversion_archivo(ruta_origen, mapa_reversion, ruta_destino=None):
    if not os.path.exists(ruta_origen) or not mapa_reversion:
        return ruta_origen

    nombre_base, ext = os.path.splitext(os.path.basename(ruta_origen))
    ext_lower = ext.lower()

    if not ruta_destino:
        target_dir = os.path.dirname(ruta_origen)
        ruta_destino = os.path.join(target_dir, f"{nombre_base}_restaurado{ext}")

    if ext_lower in [".xlsx", ".xls"]:
        wb = openpyxl.load_workbook(ruta_origen)
        for hoja in wb.worksheets:
            for fila in hoja.iter_rows():
                for celda in fila:
                    if isinstance(celda.value, str) and celda.value.strip():
                        celda.value = revertir_anonimizacion(celda.value, mapa_reversion)
        wb.save(ruta_destino)

    elif ext_lower == ".csv":
        df = pd.read_csv(ruta_origen, dtype=str)
        for columna in df.columns:
            for idx, val in enumerate(df[columna]):
                if isinstance(val, str) and val.strip():
                    df.at[idx, columna] = revertir_anonimizacion(val, mapa_reversion)
        df.to_csv(ruta_destino, index=False, encoding="utf-8")

    elif ext_lower == ".docx":
        doc = Document(ruta_origen)
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                for run in parrafo.runs:
                    if run.text.strip():
                        run.text = revertir_anonimizacion(run.text, mapa_reversion)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text.strip():
                        celda.text = revertir_anonimizacion(celda.text, mapa_reversion)
        doc.save(ruta_destino)

    elif ext_lower == ".json":
        with open(ruta_origen, "r", encoding="utf-8") as f:
            datos = json.load(f)

        def revertir_nodo(obj):
            if isinstance(obj, str):
                return revertir_anonimizacion(obj, mapa_reversion)
            elif isinstance(obj, list):
                return [revertir_nodo(item) for item in obj]
            elif isinstance(obj, dict):
                return {k: revertir_nodo(v) for k, v in obj.items()}
            else:
                return obj

        datos_restaurados = revertir_nodo(datos)
        with open(ruta_destino, "w", encoding="utf-8") as f:
            json.dump(datos_restaurados, f, indent=4, ensure_ascii=False)

    else:
        with open(ruta_origen, "r", encoding="utf-8", errors="ignore") as f:
            contenido = f.read()
        restaurado = revertir_anonimizacion(contenido, mapa_reversion)
        with open(ruta_destino, "w", encoding="utf-8") as f:
            f.write(restaurado)

    return ruta_destino


def procesar_archivo_texto(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    with open(ruta_origen, "r", encoding="utf-8") as f:
        contenido = f.read()
    
    contenido_limpio, _ = anonimizar_texto_con_ia_y_key(contenido, nombre_archivo, carpeta_salida)
    
    with open(ruta_destino, "w", encoding="utf-8") as f:
        f.write(contenido_limpio)


def procesar_markdown(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    procesar_archivo_texto(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida)


def procesar_json(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    with open(ruta_origen, "r", encoding="utf-8") as f:
        datos = json.load(f)

    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    key_map = {}
    contadores = {}

    def anonimizar_nodo(obj):
        if isinstance(obj, str):
            res, _ = anonimizar_texto_con_ia_y_key(obj, contadores_existentes=contadores, key_map_existente=key_map)
            return res
        elif isinstance(obj, list):
            return [anonimizar_nodo(item) for item in obj]
        elif isinstance(obj, dict):
            nuevo_dict = {}
            for k, v in obj.items():
                nuevo_dict[k] = anonimizar_nodo(v)
            return nuevo_dict
        else:
            return obj

    datos_limpios = anonimizar_nodo(datos)

    with open(ruta_destino, "w", encoding="utf-8") as f:
        json.dump(datos_limpios, f, indent=4, ensure_ascii=False)

    if key_map:
        guardar_llave_key(target_salida, nombre_archivo, key_map)


def procesar_dataframe_tabular(df, contadores=None, key_map=None):
    """Procedimiento Tabular por Columna unificado para CSV y Excel."""
    key_map = key_map if key_map is not None else {}
    contadores = contadores if contadores is not None else {}

    COLUMNAS_PERSONA = {"nombre", "name", "contacto", "persona", "responsable", "propietario", "representante", "entrevistado", "entrevistada"}
    COLUMNAS_EMPRESA = {"empresa", "company", "organización", "organizacion", "entidad", "razón social", "razon social", "cliente"}
    COLUMNAS_EMAIL = {"email", "correo", "e-mail", "mail"}
    COLUMNAS_URL = {"url", "web", "sitio", "website", "enlace", "link", "dominio"}
    COLUMNAS_TELEFONO = {"telefono", "teléfono", "phone", "tel", "móvil", "movil", "celular"}
    COLUMNAS_UBICACION = {"ubicacion", "ubicación", "location", "ciudad", "city", "país", "pais", "dirección", "direccion", "sede"}

    def tipo_columna_sensible(col_name):
        cl = str(col_name).lower().strip()
        cl_base = re.sub(r'\.\d+$', '', cl)
        for kw in COLUMNAS_PERSONA:
            if kw in cl_base:
                return "PERSONA"
        for kw in COLUMNAS_EMPRESA:
            if kw in cl_base:
                return "ORGANIZACION"
        for kw in COLUMNAS_EMAIL:
            if kw in cl_base:
                return "CORREO"
        for kw in COLUMNAS_URL:
            if kw in cl_base:
                return "DOMINIO"
        for kw in COLUMNAS_TELEFONO:
            if kw in cl_base:
                return "TELEFONO"
        for kw in COLUMNAS_UBICACION:
            if kw in cl_base:
                return "UBICACION"
        return None

    valor_a_etiqueta = {v: k for k, v in key_map.items()}

    for columna in df.columns:
        tipo_forzado = tipo_columna_sensible(columna)

        for idx, val in enumerate(df[columna]):
            if not isinstance(val, str) or not val.strip():
                continue

            val_limpio = val.strip()

            if tipo_forzado:
                if val_limpio.lower() in SOFTWARE_ALLOWLIST or len(val_limpio) < 2:
                    continue
                if val_limpio in valor_a_etiqueta:
                    df.at[idx, columna] = valor_a_etiqueta[val_limpio]
                else:
                    tag = tipo_forzado
                    num = contadores.get(tag, 1)
                    etiqueta = f"[{tag}_{num}]"
                    contadores[tag] = num + 1
                    key_map[etiqueta] = val_limpio
                    valor_a_etiqueta[val_limpio] = etiqueta
                    df.at[idx, columna] = etiqueta
            else:
                limpio, _ = anonimizar_texto_con_ia_y_key(val, contadores_existentes=contadores, key_map_existente=key_map)
                df.at[idx, columna] = limpio
                for k, v in key_map.items():
                    if v not in valor_a_etiqueta:
                        valor_a_etiqueta[v] = k

    return df, key_map, contadores


def procesar_tabla_csv(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    df = pd.read_csv(ruta_origen, dtype=str)
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    df, key_map, _ = procesar_dataframe_tabular(df)
    df.to_csv(ruta_destino, index=False, encoding="utf-8")
    if key_map:
        guardar_llave_key(target_salida, nombre_archivo, key_map)


def procesar_documento_word(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    doc = Document(ruta_origen)
    key_map = {}
    contadores = {}
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    
    # Procesar párrafos completos para mantener contexto semántico de IA y Regex
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            limpio, _ = anonimizar_texto_con_ia_y_key(parrafo.text, contadores_existentes=contadores, key_map_existente=key_map)
            if parrafo.runs:
                parrafo.runs[0].text = limpio
                for run in parrafo.runs[1:]:
                    run.text = ""
                    
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    limpio, _ = anonimizar_texto_con_ia_y_key(celda.text, contadores_existentes=contadores, key_map_existente=key_map)
                    celda.text = limpio
                    
    doc.save(ruta_destino)
    if key_map:
        guardar_llave_key(target_salida, nombre_archivo, key_map)


def procesar_excel(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    wb = openpyxl.load_workbook(ruta_origen)
    key_map = {}
    contadores = {}

    for hoja in wb.worksheets:
        datos = list(hoja.values)
        if not datos:
            continue
        headers = [str(h) if h is not None else f"col_{i}" for i, h in enumerate(datos[0])]
        rows = datos[1:]
        df = pd.DataFrame(rows, columns=headers).astype(str)
        df, key_map, contadores = procesar_dataframe_tabular(df, contadores=contadores, key_map=key_map)

        for col_idx, col_name in enumerate(headers, start=1):
            for row_idx, val in enumerate(df[col_name], start=2):
                celda = hoja.cell(row=row_idx, column=col_idx)
                if isinstance(celda.value, str) and celda.value.strip():
                    celda.value = val

    wb.save(ruta_destino)
    if key_map:
        guardar_llave_key(target_salida, nombre_archivo, key_map)


def procesar_pdf(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    if not fitz:
        shutil.copy(ruta_origen, ruta_destino)
        return
        
    doc = fitz.open(ruta_origen)
    mapas_acumulados = {}
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    
    for num_pag, pagina in enumerate(doc):
        texto_pagina = pagina.get_text()
        if texto_pagina.strip():
            _, mapa_pag = anonimizar_texto_con_ia_y_key(texto_pagina)
            for etiqueta, val_original in mapa_pag.items():
                if len(val_original.strip()) > 1 and val_original.lower().strip() not in SOFTWARE_ALLOWLIST:
                    rects = pagina.search_for(val_original)
                    for rect in rects:
                        pagina.add_redact_annot(rect, text=etiqueta, fill=(0.12, 0.16, 0.22), text_color=(1, 1, 1), fontsize=8)
                    mapas_acumulados[etiqueta] = val_original
                    
        pagina.apply_redactions()
        
        # C) Si no tiene texto, aplicar OCR si está disponible
        if not texto_pagina.strip() and lector_ocr:
            pix = pagina.get_pixmap(dpi=200)
            ruta_tmp = f"tmp_pag_{num_pag}.png"
            pix.save(ruta_tmp)
            procesar_imagen_ocr(ruta_tmp, ruta_tmp, f"pag_{num_pag}_temp", carpeta_salida=target_salida)
            pagina.insert_image(pagina.rect, filename=ruta_tmp)
            if os.path.exists(ruta_tmp):
                os.remove(ruta_tmp)
                    
    doc.save(ruta_destino, garbage=4, deflate=True)
    doc.close()
    
    if mapas_acumulados:
        ruta_llave = os.path.join(target_salida, f"{nombre_archivo}.reverse.key")
        with open(ruta_llave, "w", encoding="utf-8") as f:
            json.dump(mapas_acumulados, f, indent=4, ensure_ascii=False)


def procesar_imagen_ocr(ruta_origen, ruta_destino, nombre_archivo, carpeta_salida=None):
    if not cv2 or not lector_ocr:
        shutil.copy(ruta_origen, ruta_destino)
        return
        
    resultados_ocr = lector_ocr.readtext(ruta_origen)
    imagen_cv = cv2.imread(ruta_origen)
    mapas_acumulados = {}
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    terminos_corp = cargar_diccionario_corporativo()

    for (coordenadas, texto_detectado, probabilidad) in resultados_ocr:
        if len(texto_detectado.strip()) > 1:
            analisis = analyzer.analyze(
                text=texto_detectado, language="es",
                entities=["PERSON", "LOCATION", "ORGANIZATION", "EMAIL_ADDRESS", "PHONE_NUMBER", "CREDIT_CARD"]
            ) if analyzer else []
            
            dict_match = [t for t in terminos_corp if t.lower() in texto_detectado.lower()]
            
            if analisis or dict_match:
                x_min = int(min([p[0] for p in coordenadas]))
                y_min = int(min([p[1] for p in coordenadas]))
                x_max = int(max([p[0] for p in coordenadas]))
                y_max = int(max([p[1] for p in coordenadas]))
                
                label = "[CONFIDENCIAL]" if dict_match else MAPA_ETIQUETAS.get(analisis[0].entity_type, "[PII]")
                
                cv2.rectangle(imagen_cv, (x_min, y_min), (x_max, y_max), (30, 24, 18), -1)
                cv2.rectangle(imagen_cv, (x_min, y_min), (x_max, y_max), (248, 189, 56), 1)
                
                h_box = max(1, y_max - y_min)
                scale = max(0.3, min(0.5, h_box / 35.0))
                cv2.putText(imagen_cv, label, (x_min + 3, y_min + int(h_box * 0.7)), cv2.FONT_HERSHEY_SIMPLEX, scale, (255, 255, 255), 1, cv2.LINE_AA)
                mapas_acumulados[f"[IMAGE_PII_{len(mapas_acumulados)+1}]"] = texto_detectado

    cv2.imwrite(ruta_destino, imagen_cv)
    if mapas_acumulados:
        ruta_llave = os.path.join(target_salida, f"{nombre_archivo}.reverse.key")
        with open(ruta_llave, "w", encoding="utf-8") as f:
            json.dump(mapas_acumulados, f, indent=4, ensure_ascii=False)


def ejecutar_procesamiento_lotes(carpeta_salida=None):
    target_salida = carpeta_salida or CARPETA_SALIDA_DEFECTO
    os.makedirs(target_salida, exist_ok=True)
    os.makedirs(CARPETA_PROCESADOS, exist_ok=True)

    archivos = os.listdir(CARPETA_ENTRADA)
    archivos = [f for f in archivos if not f.startswith("~$") and os.path.isfile(os.path.join(CARPETA_ENTRADA, f))]
    
    if not archivos:
        print(f"📁 La carpeta '{CARPETA_ENTRADA}' está vacía. Agregue archivos para procesar.")
        return 0

    print(f"🚀 LIA VAULT: Procesando lote de {len(archivos)} archivos -> '{target_salida}'...")
    procesados_exitosos = 0
    
    for nombre_archivo in archivos:
        ruta_origen = os.path.join(CARPETA_ENTRADA, nombre_archivo)
        nombre_sanitizado, fn_map = sanitizar_nombre_archivo(nombre_archivo, carpeta_salida=target_salida)
        ruta_destino = os.path.join(target_salida, nombre_sanitizado)
        ext = nombre_archivo.lower()
        
        try:
            if ext.endswith(".txt"):
                procesar_archivo_texto(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".md"):
                procesar_markdown(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".json"):
                procesar_json(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".csv"):
                procesar_tabla_csv(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".docx"):
                procesar_documento_word(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".xlsx"):
                procesar_excel(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith(".pdf"):
                procesar_pdf(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            elif ext.endswith((".jpg", ".jpeg", ".png")):
                procesar_imagen_ocr(ruta_origen, ruta_destino, nombre_archivo, target_salida)
            else:
                print(f"⚠️ Formato no compatible: {nombre_archivo}")
                continue
            
            # Mover archivo original procesado exitosamente a /procesados para no volver a encriptarlo
            ruta_procesados = os.path.join(CARPETA_PROCESADOS, nombre_archivo)
            if os.path.exists(ruta_procesados):
                os.remove(ruta_procesados)
            shutil.move(ruta_origen, ruta_procesados)
            procesados_exitosos += 1
            print(f"✅ Procesado y movido a /procesados: {nombre_archivo}")

        except Exception as e:
            print(f"❌ Error al procesar {nombre_archivo}: {e}")
            
    return procesados_exitosos


def guardar_llave_key(target_salida, nombre_archivo, key_map):
    if not key_map or not nombre_archivo:
        return
    r1 = os.path.join(target_salida, f"{nombre_archivo}.key")
    r2 = os.path.join(target_salida, f"{nombre_archivo}.reverse.key")
    try:
        with open(r1, "w", encoding="utf-8") as f:
            json.dump(key_map, f, indent=4, ensure_ascii=False)
        with open(r2, "w", encoding="utf-8") as f:
            json.dump(key_map, f, indent=4, ensure_ascii=False)
    except Exception as e:
        print(f"Error guardando .key: {e}")


def desanonimizar_texto(texto_anonimizado, mapa_llave):
    """Reemplaza los tokens anonimizados por su valor original utilizando el diccionario de llave .key."""
    texto_restaurado = texto_anonimizado
    if isinstance(mapa_llave, str):
        try:
            mapa_llave = json.loads(mapa_llave)
        except Exception:
            pass
            
    if isinstance(mapa_llave, dict):
        for token, valor_original in mapa_llave.items():
            texto_restaurado = texto_restaurado.replace(token, valor_original)
    return texto_restaurado


def desanonimizar_archivo(ruta_archivo, mapa_llave, ruta_destino_restaurada):
    """Lee un archivo anonimizado y aplica la llave para restaurar sus valores originales en el destino."""
    if isinstance(mapa_llave, str):
        try:
            mapa_llave = json.loads(mapa_llave)
        except Exception:
            pass
            
    if not os.path.exists(ruta_archivo) or not isinstance(mapa_llave, dict):
        return False
        
    try:
        with open(ruta_archivo, "r", encoding="utf-8", errors="ignore") as f_in:
            txt = f_in.read()
        txt_restaurado = desanonimizar_texto(txt, mapa_llave)
        with open(ruta_destino_restaurada, "w", encoding="utf-8") as f_out:
            f_out.write(txt_restaurado)
        return True
    except Exception as e:
        print(f"Error desanonimizando archivo: {e}")
        return False


if __name__ == "__main__":
    ejecutar_procesamiento_lotes()
